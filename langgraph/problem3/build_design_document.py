from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSET_DIR = ROOT / "design_assets"
OUTPUT = ROOT / "multi_agent_research_system_design.docx"

NAVY = "173B67"
BLUE = "2867B2"
SKY = "EAF2FC"
PALE = "F5F8FC"
INK = "17212B"
MUTED = "536273"
GREEN = "287A5B"
GOLD = "C58A20"
RED = "B24A4A"
WHITE = "FFFFFF"
LIGHT_BORDER = "CBD7E6"


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    if bold:
        candidates = [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/Library/Fonts/Arial Bold.ttf",
        ] + candidates
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline=LIGHT_BORDER, radius=24, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=f"#{fill}", outline=f"#{outline}", width=width)


def centered_text(draw, xy, text, size=30, color=INK, bold=False, max_chars=22):
    x1, y1, x2, y2 = xy
    lines = []
    for part in text.split("\n"):
        lines.extend(wrap(part, max_chars) or [""])
    face = font(size, bold)
    heights = []
    widths = []
    for line in lines:
        box = draw.textbbox((0, 0), line, font=face)
        widths.append(box[2] - box[0])
        heights.append(box[3] - box[1])
    gap = 8
    total_h = sum(heights) + gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_h) / 2
    for line, line_w, line_h in zip(lines, widths, heights):
        draw.text((x1 + (x2 - x1 - line_w) / 2, y), line, fill=f"#{color}", font=face)
        y += line_h + gap


def arrow(draw, start, end, color=BLUE, width=6, head=18):
    draw.line([start, end], fill=f"#{color}", width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    base_x, base_y = x2 - head * ux, y2 - head * uy
    points = [
        (x2, y2),
        (base_x + head * 0.55 * px, base_y + head * 0.55 * py),
        (base_x - head * 0.55 * px, base_y - head * 0.55 * py),
    ]
    draw.polygon(points, fill=f"#{color}")


def double_arrow(draw, start, end, color=BLUE, width=5):
    arrow(draw, start, end, color, width)
    arrow(draw, end, start, color, width)


def title_on_canvas(draw, title, subtitle=None):
    draw.text((80, 45), title, font=font(44, True), fill=f"#{NAVY}")
    if subtitle:
        draw.text((82, 105), subtitle, font=font(23), fill=f"#{MUTED}")


def build_architecture():
    image = Image.new("RGB", (1600, 980), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title_on_canvas(draw, "Multi-Agent Research Architecture", "A coordinator manages specialists through a shared collaborative state")

    coordinator = (575, 155, 1025, 275)
    rounded_box(draw, coordinator, NAVY, NAVY)
    centered_text(draw, coordinator, "Research Coordinator", 35, WHITE, True)

    state = (460, 400, 1140, 585)
    rounded_box(draw, state, SKY, BLUE, radius=30, width=5)
    centered_text(
        draw,
        state,
        "Collaborative State\nContext | Contributions | Cross-Agent Insights\nReasoning | Synthesis Threads",
        28,
        NAVY,
        True,
        48,
    )
    double_arrow(draw, (800, 275), (800, 400), BLUE, 6)

    agents = [
        ((55, 705, 330, 865), "Web Research\nAgent"),
        ((360, 705, 635, 865), "Data Analysis\nAgent"),
        ((665, 705, 940, 865), "Trend Analysis\nAgent"),
        ((970, 705, 1245, 865), "Competitive\nIntelligence Agent"),
        ((1275, 705, 1545, 865), "Synthesis\nAgent"),
    ]
    for index, (box, label) in enumerate(agents):
        fill = BLUE if index < 4 else GREEN
        rounded_box(draw, box, fill, fill)
        centered_text(draw, box, label, 27, WHITE, True, 21)
        center_x = (box[0] + box[2]) // 2
        double_arrow(draw, (center_x, box[1]), (800, state[3]), MUTED, 4)

    path = ASSET_DIR / "architecture.png"
    image.save(path)
    return path


def build_state_flow():
    image = Image.new("RGB", (1600, 900), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title_on_canvas(draw, "Collaborative State Lifecycle", "Agents do more than deposit facts: they refine a shared understanding")

    boxes = [
        ((60, 230, 330, 390), "1. Shared\nResearch Context", NAVY),
        ((385, 230, 655, 390), "2. Agent\nContributions", BLUE),
        ((710, 230, 980, 390), "3. Cross-Agent\nInsights", BLUE),
        ((1035, 230, 1305, 390), "4. Collaborative\nReasoning", NAVY),
        ((710, 580, 980, 740), "5. Synthesis\nThreads", GREEN),
    ]
    for box, label, fill in boxes:
        rounded_box(draw, box, fill, fill)
        centered_text(draw, box, label, 29, WHITE, True, 22)
    for start, end in [
        ((330, 310), (385, 310)),
        ((655, 310), (710, 310)),
        ((980, 310), (1035, 310)),
        ((1170, 390), (980, 660)),
    ]:
        arrow(draw, start, end)

    feedback = (120, 560, 570, 750)
    rounded_box(draw, feedback, PALE, LIGHT_BORDER)
    centered_text(
        draw,
        feedback,
        "Feedback loop\nGaps, contradictions, or low confidence\ncreate new research tasks",
        27,
        INK,
        True,
        36,
    )
    arrow(draw, (710, 660), (570, 660), GOLD)
    arrow(draw, (345, 560), (195, 390), GOLD)

    note = (1100, 540, 1515, 760)
    rounded_box(draw, note, SKY, BLUE)
    centered_text(
        draw,
        note,
        "Readiness gate\nSynthesis begins only when coverage,\nconfidence, and contradiction checks pass",
        25,
        NAVY,
        True,
        38,
    )
    arrow(draw, (1035, 660), (980, 660), GREEN)

    path = ASSET_DIR / "state_lifecycle.png"
    image.save(path)
    return path


def build_dynamic_workflow():
    image = Image.new("RGB", (1600, 1120), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title_on_canvas(draw, "Adaptive Collaboration Workflow", "The route changes according to question complexity and emerging evidence")

    items = [
        ((545, 145, 1055, 250), "Receive Research Question", NAVY),
        ((545, 315, 1055, 430), "Classify Scope, Evidence Needs,\nand Complexity", BLUE),
        ((545, 495, 1055, 610), "Select Agents and Dependencies", BLUE),
        ((120, 700, 485, 825), "Parallel Track\nWeb + Data Research", SKY),
        ((615, 700, 985, 825), "Sequential Track\nResearch -> Trends -> Competition", SKY),
        ((1115, 700, 1480, 825), "Targeted Track\nActivate Only Needed Specialists", SKY),
        ((545, 930, 1055, 1040), "Synthesis Readiness Gate", GREEN),
    ]
    for box, label, fill in items:
        outline = fill if fill != SKY else BLUE
        rounded_box(draw, box, fill, outline)
        text_color = WHITE if fill not in (SKY, PALE) else NAVY
        centered_text(draw, box, label, 29, text_color, True, 34)

    arrow(draw, (800, 250), (800, 315))
    arrow(draw, (800, 430), (800, 495))
    arrow(draw, (680, 610), (300, 700))
    arrow(draw, (800, 610), (800, 700))
    arrow(draw, (920, 610), (1300, 700))
    arrow(draw, (300, 825), (670, 930))
    arrow(draw, (800, 825), (800, 930))
    arrow(draw, (1300, 825), (930, 930))

    draw.text((60, 872), "Independent evidence", font=font(22, True), fill=f"#{MUTED}")
    draw.text((666, 872), "Dependency-driven", font=font(22, True), fill=f"#{MUTED}")
    draw.text((1220, 872), "Narrow request", font=font(22, True), fill=f"#{MUTED}")

    path = ASSET_DIR / "dynamic_workflow.png"
    image.save(path)
    return path


def build_interaction_example():
    image = Image.new("RGB", (1600, 1050), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title_on_canvas(draw, "Example: Electric Vehicle Market Analysis", "A finding becomes valuable when another agent tests, connects, and explains it")

    boxes = [
        ((70, 175, 465, 310), "Web Research\nPolicy incentives are expanding", BLUE),
        ((600, 175, 995, 310), "Data Analysis\nSales growth is uneven by region", BLUE),
        ((1130, 175, 1525, 310), "Trend Analysis\nAdoption may accelerate in 2027", NAVY),
        ((1130, 475, 1525, 610), "Competitive Intelligence\nCompetitors shift investment to charging", NAVY),
        ((600, 745, 995, 900), "Synthesis Agent\nConnect policy, demand, and strategy", GREEN),
        ((70, 745, 465, 900), "Final Insight\nInfrastructure, not demand alone,\nbecomes the strategic constraint", GREEN),
    ]
    for box, label, fill in boxes:
        rounded_box(draw, box, fill, fill)
        centered_text(draw, box, label, 24, WHITE, True, 24)

    arrow(draw, (465, 240), (600, 240))
    arrow(draw, (995, 240), (1130, 240))
    arrow(draw, (1327, 310), (1327, 475))
    arrow(draw, (1130, 542), (925, 745))
    arrow(draw, (797, 310), (797, 745))
    arrow(draw, (600, 822), (465, 822), GREEN)

    feedback = (95, 455, 800, 600)
    rounded_box(draw, feedback, PALE, GOLD)
    centered_text(
        draw,
        feedback,
        "Cross-agent challenge:\nDoes charging capacity explain the regional sales gap?\nTrigger targeted infrastructure research.",
        27,
        INK,
        True,
        54,
    )
    arrow(draw, (800, 525), (1130, 525), GOLD)
    arrow(draw, (250, 455), (250, 310), GOLD)

    path = ASSET_DIR / "interaction_example.png"
    image.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run(r2)
    else:
        set_run(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item))


def add_compact_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        set_run(p.add_run(item), size=10)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    style_num_id = int(doc.styles["List Number"].element.pPr.numPr.numId.val)
    source_num = numbering.find(f".//{{{numbering.nsmap['w']}}}num[@{{{numbering.nsmap['w']}}}numId='{style_num_id}']")
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_num_id = max(existing_ids) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    new_num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)

    for item in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
        num_pr.get_or_add_numId().set(qn("w:val"), str(new_num_id))
        set_run(p.add_run(item))


def add_callout(doc, label, text, fill=SKY, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label + "  "), bold=True, color=accent)
    set_run(p.add_run(text), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(image_path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    set_run(cap.add_run(caption), size=9, italic=True, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_page_break(doc):
    doc.add_page_break()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 17, NAVY, 15, 7),
        ("Heading 2", 13.5, BLUE, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.15

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(hp.add_run("Multi-Agent Research System | Design Document"), size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(fp.add_run("Problem 3 | Designing a Multi-Agent Research System"), size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(92)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run("DESIGN DOCUMENT"), size=11, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(10)
    set_run(title.add_run("Designing a Multi-Agent\nResearch System"), size=30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_run(
        subtitle.add_run("Collaborative intelligence for comprehensive market analysis"),
        size=14,
        color=BLUE,
    )

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(line.add_run("Five specialized agents. One evolving shared understanding."), size=11.5, italic=True, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.6, 3.8])
    data = [
        ("System", "Multi-Agent Market Research Assistant"),
        ("Domain", "Consulting and Market Intelligence"),
        ("Design Focus", "State, coordination, adaptive workflow, synthesis"),
        ("Output", "Evidence-based market analysis report"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell_shading(row.cells[0], SKY)
        for cell in row.cells:
            set_cell_margins(cell, 95, 130, 95, 130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_run(row.cells[0].paragraphs[0].add_run(label), bold=True, color=NAVY)
        set_run(row.cells[1].paragraphs[0].add_run(value), color=INK)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    set_run(p.add_run("Prepared from the original Problem 3 assignment brief"), size=9.5, color=MUTED)
    add_page_break(doc)


def build_document():
    ASSET_DIR.mkdir(exist_ok=True)
    figures = {
        "architecture": build_architecture(),
        "state": build_state_flow(),
        "workflow": build_dynamic_workflow(),
        "example": build_interaction_example(),
    }

    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "This design proposes a coordinated multi-agent research system for producing comprehensive market analysis reports. "
        "A research coordinator interprets the client question, activates the right specialist agents, manages dependencies, "
        "and decides when the evidence is mature enough for synthesis."
    )
    add_body(
        doc,
        "The design does not treat the agents as isolated report writers. Every agent reads from and contributes to a structured "
        "collaborative state. Findings can therefore be challenged, extended, connected, or converted into new research tasks. "
        "The final report is built from synthesis threads that preserve the relationships among evidence, quantitative analysis, "
        "market trends, and competitor behavior."
    )
    add_callout(
        doc,
        "Core design principle",
        "The system creates value when one agent's work changes what another agent investigates or concludes.",
    )

    add_heading(doc, "Design Goals", 2)
    add_bullets(
        doc,
        [
            "Create shared understanding instead of merely collecting separate agent outputs.",
            "Make dependencies and handoffs explicit so agents receive the context they need.",
            "Adapt the workflow to the question rather than running every agent in the same order.",
            "Support investigation loops when evidence is incomplete, contradictory, or low confidence.",
            "Begin synthesis only after coverage and quality conditions are satisfied.",
        ],
    )

    add_heading(doc, "1. System Architecture", 1)
    add_body(
        doc,
        "The architecture uses a coordinator pattern around a collaborative state. The coordinator has visibility into all "
        "specialists and the shared state, while specialists communicate through structured contributions and events. This "
        "reduces uncontrolled agent-to-agent links and makes the research process observable."
    )
    add_figure(
        doc,
        figures["architecture"],
        "Figure 1. The coordinator selects and sequences specialists; the collaborative state carries shared knowledge.",
    )

    add_heading(doc, "Architectural Responsibilities", 2)
    architecture_rows = [
        ("Research Coordinator", "Interprets the request, creates the plan, activates agents, resolves dependencies, and applies readiness gates."),
        ("Specialist Agents", "Perform bounded research or analysis using the context and evidence relevant to their expertise."),
        ("Collaborative State", "Stores research context, evidence, reasoning, links, unresolved questions, and synthesis threads."),
        ("Synthesis Agent", "Builds an integrated report from validated threads rather than concatenating agent summaries."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.75, 4.75])
    for i, text in enumerate(("Component", "Responsibility")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), bold=True, color=WHITE)
    for label, detail in architecture_rows:
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(label), bold=True, color=NAVY)
        set_run(cells[1].paragraphs[0].add_run(detail))
        for cell in cells:
            set_cell_margins(cell)

    add_page_break(doc)
    add_heading(doc, "2. Collaborative State Architecture", 1)
    add_body(
        doc,
        "The shared state is organized into five layers. Each layer has a distinct purpose, owner expectations, and quality "
        "rules. This prevents the state from becoming a single unstructured container that every agent edits differently."
    )
    add_figure(
        doc,
        figures["state"],
        "Figure 2. Contributions are converted into connected reasoning and synthesis threads, with feedback when gaps remain.",
    )

    add_page_break(doc)
    state_fields = [
        ("Shared Research Context", "research_question, client_objective, scope, geography, time_horizon, definitions, constraints, success_criteria", "Coordinator", "Stable framing used by every agent"),
        ("Agent Contributions", "agent_id, finding, evidence, source, confidence, assumptions, timestamp, related_question", "Specialist agents", "Traceable domain findings"),
        ("Cross-Agent Insights", "linked_findings, relationship_type, explanation, supporting_agents, confidence", "Any agent or coordinator", "Connections across domains"),
        ("Collaborative Reasoning", "hypotheses, confirmations, contradictions, open_questions, agreed_interpretations", "Coordinator and specialists", "Evolving shared understanding"),
        ("Synthesis Threads", "thread_title, claim, supporting_evidence, implications, risks, unresolved_items", "Synthesis agent", "Report-ready narratives"),
    ]
    add_heading(doc, "State Schema", 2)
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.25, 2.55, 1.05, 1.65])
    headers = ("State Layer", "Representative Fields", "Primary Writer", "Purpose")
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.5, bold=True, color=WHITE)
    for layer, fields, writer, purpose in state_fields:
        cells = table.add_row().cells
        values = (layer, fields, writer, purpose)
        for i, value in enumerate(values):
            set_run(cells[i].paragraphs[0].add_run(value), size=8.8, bold=(i == 0), color=NAVY if i == 0 else INK)
            set_cell_margins(cells[i], 90, 95, 90, 95)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_heading(doc, "State Update Rules", 2)
    add_numbered(
        doc,
        [
            "Append evidence rather than silently replacing another agent's finding.",
            "Attach a source and confidence level to every material claim.",
            "Represent disagreement explicitly as a contradiction or competing hypothesis.",
            "Link new insights to the findings that support them.",
            "Promote content into synthesis threads only after evidence and dependency checks.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "3. Agent Coordination Matrix", 1)
    add_body(
        doc,
        "Each agent has a bounded role, but its work is intentionally connected to the others. The matrix below defines what "
        "each agent needs, contributes, responds to, and signals."
    )
    matrix = [
        (
            "Web Research",
            "Question, scope, keywords, unresolved evidence gaps",
            "Articles, reports, facts, citations, source quality",
            "New task; missing evidence; claim needing verification",
            "Source found; weak source; conflicting external evidence",
        ),
        (
            "Data Analysis",
            "Datasets, definitions, web findings, metrics to test",
            "Calculations, comparisons, charts, anomalies, confidence",
            "New dataset; quantitative claim; regional or temporal gap",
            "Statistical pattern; outlier; insufficient data",
        ),
        (
            "Trend Analysis",
            "Research evidence, data results, time horizon",
            "Drivers, trajectories, scenarios, projections, uncertainty",
            "Stable baseline evidence; material change signal",
            "Emerging trend; turning point; projection risk",
        ),
        (
            "Competitive Intelligence",
            "Market trends, company evidence, market metrics",
            "Competitor moves, positioning, capability gaps, implications",
            "Competitor event; trend with strategic impact",
            "Strategic shift; market response; competitor contradiction",
        ),
        (
            "Synthesis",
            "Validated findings, cross-agent insights, open issues, threads",
            "Integrated narrative, implications, recommendations, report",
            "Coverage threshold met; major contradictions resolved",
            "Draft ready; evidence gap; final review required",
        ),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.05, 1.35, 1.55, 1.3, 1.25])
    for i, text in enumerate(("Agent", "Needs", "Contributes", "Responds To", "Signals")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), size=8.7, bold=True, color=WHITE)
    for row_data in matrix:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            if i == 0:
                set_cell_shading(cells[i], SKY)
            set_run(cells[i].paragraphs[0].add_run(value), size=8.15, bold=(i == 0), color=NAVY if i == 0 else INK)
            set_cell_margins(cells[i], 80, 85, 80, 85)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    add_heading(doc, "Coordination Pattern", 2)
    add_bullets(
        doc,
        [
            "The coordinator creates tasks with explicit inputs, expected outputs, and dependencies.",
            "Agents publish structured contributions to the shared state rather than sending informal messages.",
            "Important findings generate events that can activate another specialist.",
            "The coordinator tracks open questions, coverage, confidence, and contradictions.",
            "The synthesis agent remains downstream of evidence maturity, not simply elapsed time.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "4. Dynamic Workflow Logic", 1)
    add_body(
        doc,
        "The system first classifies the research request, then chooses an execution pattern. A narrow request can use a targeted "
        "path; independent evidence gathering can run in parallel; dependent reasoning is sequenced so later agents inherit "
        "validated context."
    )
    add_figure(
        doc,
        figures["workflow"],
        "Figure 3. The coordinator selects parallel, sequential, or targeted collaboration and converges at a synthesis gate.",
    )

    add_heading(doc, "Agent Activation Framework", 2)
    activation = [
        ("Factual or source-led question", "Web Research", "Data Analysis if quantitative validation is required"),
        ("Market size or performance question", "Web Research + Data Analysis in parallel", "Trend Analysis after baseline validation"),
        ("Future outlook question", "Web Research + Data Analysis", "Trend Analysis, then Competitive Intelligence"),
        ("Competitor strategy question", "Web Research + Competitive Intelligence", "Data and Trend agents when market context is needed"),
        ("Comprehensive market report", "All evidence agents", "Trend and Competitive agents build on research/data; Synthesis runs last"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [2.0, 2.1, 2.4])
    for i, text in enumerate(("Question Type", "Initial Activation", "Dependency Logic")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.3, bold=True, color=WHITE)
    for row_data in activation:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_run(cells[i].paragraphs[0].add_run(value), size=8.8, bold=(i == 0), color=NAVY if i == 0 else INK)
            set_cell_margins(cells[i])

    add_heading(doc, "Parallel vs. Sequential Decision", 2)
    add_bullets(
        doc,
        [
            "Run in parallel when agents can collect independent evidence from the same stable research context.",
            "Run sequentially when an agent's task definition depends on another agent's result.",
            "Use a hybrid workflow when broad evidence collection is parallel but interpretation requires ordered handoffs.",
            "Pause or backtrack when a critical input is absent, confidence is low, or findings materially conflict.",
        ],
    )

    add_heading(doc, "Synthesis Readiness Conditions", 2)
    add_callout(
        doc,
        "Readiness gate",
        "Begin synthesis when required domains are covered, major claims have evidence, confidence is acceptable, critical contradictions are resolved or disclosed, and no blocking research task remains.",
        fill="EAF5F0",
        accent=GREEN,
    )

    add_heading(doc, "5. Investigation, Backtracking, and Quality Control", 1)
    add_body(
        doc,
        "Collaboration is not complete after the first pass. Agents may discover evidence gaps or contradictions that require "
        "additional work. The coordinator turns these discoveries into explicit follow-up tasks and preserves both the original "
        "finding and the resolution."
    )

    quality_rows = [
        ("Missing evidence", "A material claim has no reliable source", "Web Research receives a targeted verification task"),
        ("Quantitative inconsistency", "Reported figures disagree across sources", "Data Analysis reconciles definitions, dates, and units"),
        ("Low-confidence projection", "Trend conclusion depends on weak assumptions", "Trend Analysis creates scenarios and records uncertainty"),
        ("Competitive contradiction", "Observed company action conflicts with stated strategy", "Competitive Intelligence investigates behavior and timing"),
        ("Narrative gap", "A report thread lacks evidence or implication", "Synthesis returns the thread to the relevant specialist"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1.65, 2.35, 2.5])
    for i, text in enumerate(("Condition", "Detection", "Backtracking Action")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9.3, bold=True, color=WHITE)
    for row_data in quality_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_run(cells[i].paragraphs[0].add_run(value), size=9, bold=(i == 0), color=RED if i == 0 else INK)
            set_cell_margins(cells[i])

    add_heading(doc, "Completion Controls", 2)
    add_compact_bullets(
        doc,
        [
            "Coverage: every required report domain has at least one validated contribution.",
            "Traceability: important claims link to sources and originating agent contributions.",
            "Consistency: contradictions are resolved, bounded, or transparently reported.",
            "Confidence: uncertainty is recorded and reflected in conclusions.",
            "Coherence: synthesis threads connect evidence to implications rather than listing facts.",
        ],
    )

    add_heading(doc, "6. Worked Collaboration Example", 1)
    add_body(
        doc,
        "Consider the research question: “How will the electric vehicle market evolve over the next three years, and where should "
        "a charging-network provider invest?” The example below shows how an isolated observation becomes a stronger strategic "
        "insight through collaboration."
    )
    add_figure(
        doc,
        figures["example"],
        "Figure 4. Cross-agent reasoning converts separate findings into a defensible strategic conclusion.",
    )
    add_heading(doc, "Example Reasoning Trace", 2)
    add_numbered(
        doc,
        [
            "Web Research identifies new incentives and infrastructure policy.",
            "Data Analysis finds that sales growth varies sharply across regions.",
            "Trend Analysis proposes that adoption will accelerate but records infrastructure as a limiting assumption.",
            "Competitive Intelligence observes competitors redirecting investment toward charging capacity.",
            "The coordinator creates a targeted question linking charging coverage to regional sales performance.",
            "The synthesis agent forms a report thread: infrastructure availability may be the strategic constraint that determines where demand converts into adoption.",
        ],
    )

    add_heading(doc, "Why This Is Collaborative Intelligence", 2)
    add_body(
        doc,
        "No single agent produced the final conclusion. Web evidence supplied policy context, quantitative analysis exposed regional "
        "variation, trend reasoning framed the future trajectory, and competitive analysis revealed how firms were responding. "
        "The shared state made these relationships visible, and the feedback loop prompted a new investigation. The final insight "
        "therefore emerged from connected reasoning rather than aggregation."
    )

    add_heading(doc, "7. Report Assembly", 1)
    add_body(
        doc,
        "The synthesis agent assembles the final deliverable from validated synthesis threads. Each section should preserve the "
        "relationship between evidence, interpretation, business implication, and uncertainty."
    )
    report_sections = [
        ("Executive Summary", "Most important conclusions, implications, and recommended actions"),
        ("Market Overview", "Scope, size, segments, definitions, and baseline conditions"),
        ("Data Insights", "Quantitative findings, comparisons, anomalies, and limitations"),
        ("Trends and Projections", "Drivers, scenarios, time horizon, and uncertainty"),
        ("Competitive Landscape", "Competitor moves, positioning, responses, and strategic gaps"),
        ("Integrated Implications", "Cross-agent synthesis threads connecting market evidence to decisions"),
        ("Conclusion", "Priorities, risks, open questions, and next research steps"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [2.0, 4.5])
    for i, text in enumerate(("Report Section", "Content")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), bold=True, color=WHITE)
    for section, content in report_sections:
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(section), bold=True, color=NAVY)
        set_run(cells[1].paragraphs[0].add_run(content))
        for cell in cells:
            set_cell_margins(cell)

    add_heading(doc, "8. Collaboration Rationale", 1)
    add_body(
        doc,
        "The proposed design satisfies the assignment's requirement for true collaborative intelligence in four ways."
    )
    rationale = [
        ("Shared context creates alignment.", "All agents interpret the same objective, scope, definitions, and constraints."),
        ("Structured contributions create traceability.", "Findings retain evidence, confidence, assumptions, and ownership."),
        ("Cross-agent insights create new knowledge.", "The system records relationships that are not present in any isolated contribution."),
        ("Adaptive coordination improves relevance.", "Agents are activated, sequenced, or re-engaged according to the question and emerging findings."),
    ]
    for title, detail in rationale:
        add_body(doc, f"{title} {detail}", bold_lead=title)

    add_callout(
        doc,
        "Final design outcome",
        "The system produces a report whose conclusions are supported by connected evidence and collective reasoning, not a bundle of independent summaries.",
        fill="EAF5F0",
        accent=GREEN,
    )

    add_heading(doc, "9. Completion Criteria Mapping", 1)
    mapping = [
        ("Collaborative state enables intelligence beyond data sharing", "Five-layer state architecture with explicit reasoning and synthesis layers"),
        ("Coordination patterns show agents building on each other", "Coordination matrix, dependencies, triggers, and event signals"),
        ("Workflow adapts to research complexity", "Parallel, sequential, targeted, and hybrid activation framework"),
        ("Rationale explains emergent intelligence", "Worked example and collaboration rationale"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [3.0, 3.5])
    for i, text in enumerate(("Assignment Criterion", "Where Addressed")):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        set_run(table.rows[0].cells[i].paragraphs[0].add_run(text), bold=True, color=WHITE)
    for criterion, location in mapping:
        cells = table.add_row().cells
        set_run(cells[0].paragraphs[0].add_run(criterion), bold=True, color=NAVY)
        set_run(cells[1].paragraphs[0].add_run(location))
        for cell in cells:
            set_cell_margins(cell)

    doc.core_properties.title = "Designing a Multi-Agent Research System"
    doc.core_properties.subject = "Collaborative multi-agent architecture and workflow design"
    doc.core_properties.author = "LLM Concepts"
    doc.core_properties.keywords = "multi-agent, collaborative state, market research, workflow, synthesis"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
